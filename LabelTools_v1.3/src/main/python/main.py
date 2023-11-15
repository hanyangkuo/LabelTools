from fbs_runtime.application_context.PyQt5 import ApplicationContext
import pkg_resources.py2_warn
import xlrd 
import pandas as pd
import numpy as np
import xlwings as xw
import os, sys, glob
from xlwings.utils import rgb_to_int
from enum import Enum
import docx
from docx import Document
from calendar import monthrange
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import QStringListModel 

is_label = False
head_1 = ["日期", "單位", "員工"]
head_2 = [
    "日期", "夜觀", "公差", "公假", "事假", "病假", "喪假",
    "備勤", "其他類別（一）夜觀、公差、公假、事假、病假、喪假, 日期",
    "災防", "公傷", "生理假", "產前假", "婚假", "分娩", "流產假", "",
    "其他類別（二）災防、公傷、生理假、產前假、婚假、分娩、公出、看護"
]

if not os.path.exists("教保科人員.txt"):
    f = open("教保科人員.txt",'w')
    f.close()

if not os.path.exists("教保科役男.txt"):
    f = open("教保科役男.txt",'w')
    f.close()

class Departments(Enum):
    博 = 0
    仁 = 1
    慈 = 2
    永 = 3
    弘 = 4
    崇 = 5
    承 = 6
    信 = 7
    友 = 8
    祥 = 9
    教 = 10
    
class Color(Enum):
    早班 = (0, 204, 255)
    假日班 = (220, 124, 213)
    中班 = (255, 255, 0)
    中班B = (204, 153, 0)
    夜班 = (192, 192, 192)
    輪休 = (204, 255, 204)
    上休 = xw.utils.rgb_to_int((255, 0, 255))
    下休 = xw.utils.rgb_to_int((0, 0, 255))
    備勤 = (255, 51, 153)

class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(671, 548)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.scrollArea = QtWidgets.QScrollArea(self.centralwidget)
        self.scrollArea.setGeometry(QtCore.QRect(11, 11, 651, 51))
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setObjectName("scrollArea")
        self.scrollAreaWidgetContents = QtWidgets.QWidget()
        self.scrollAreaWidgetContents.setGeometry(QtCore.QRect(0, 0, 649, 49))
        self.scrollAreaWidgetContents.setObjectName("scrollAreaWidgetContents")
        self.gridLayout_2 = QtWidgets.QGridLayout(self.scrollAreaWidgetContents)
        self.gridLayout_2.setObjectName("gridLayout_2")
        self.label_Excel = QtWidgets.QLabel(self.scrollAreaWidgetContents)
        font = QtGui.QFont()
        font.setFamily("標楷體")
        font.setPointSize(12)
        self.label_Excel.setFont(font)
        self.label_Excel.setObjectName("label_Excel")
        self.gridLayout_2.addWidget(self.label_Excel, 0, 0, 1, 1)
        self.lineEdit_Excel = QtWidgets.QLineEdit(self.scrollAreaWidgetContents)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.lineEdit_Excel.setFont(font)
        self.lineEdit_Excel.setObjectName("lineEdit_Excel")
        self.gridLayout_2.addWidget(self.lineEdit_Excel, 0, 1, 1, 1)
        self.scrollArea.setWidget(self.scrollAreaWidgetContents)
        self.scrollArea_2 = QtWidgets.QScrollArea(self.centralwidget)
        self.scrollArea_2.setGeometry(QtCore.QRect(11, 68, 651, 87))
        self.scrollArea_2.setWidgetResizable(True)
        self.scrollArea_2.setObjectName("scrollArea_2")
        self.scrollAreaWidgetContents_2 = QtWidgets.QWidget()
        self.scrollAreaWidgetContents_2.setGeometry(QtCore.QRect(0, 0, 649, 85))
        self.scrollAreaWidgetContents_2.setObjectName("scrollAreaWidgetContents_2")
        self.gridLayout = QtWidgets.QGridLayout(self.scrollAreaWidgetContents_2)
        self.gridLayout.setObjectName("gridLayout")
        self.label_Serviceman = QtWidgets.QLabel(self.scrollAreaWidgetContents_2)
        font = QtGui.QFont()
        font.setFamily("標楷體")
        font.setPointSize(12)
        self.label_Serviceman.setFont(font)
        self.label_Serviceman.setObjectName("label_Serviceman")
        self.gridLayout.addWidget(self.label_Serviceman, 0, 0, 1, 1)
        self.lineEdit_Serviceman = QtWidgets.QLineEdit(self.scrollAreaWidgetContents_2)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.lineEdit_Serviceman.setFont(font)
        self.lineEdit_Serviceman.setObjectName("lineEdit_Serviceman")
        self.gridLayout.addWidget(self.lineEdit_Serviceman, 0, 1, 1, 1)
        self.pushButton_Serviceman = QtWidgets.QPushButton(self.scrollAreaWidgetContents_2)
        self.pushButton_Serviceman.setObjectName("pushButton_Serviceman")
        self.gridLayout.addWidget(self.pushButton_Serviceman, 0, 2, 1, 1)
        self.label_Schedule = QtWidgets.QLabel(self.scrollAreaWidgetContents_2)
        font = QtGui.QFont()
        font.setFamily("標楷體")
        font.setPointSize(12)
        self.label_Schedule.setFont(font)
        self.label_Schedule.setObjectName("label_Schedule")
        self.gridLayout.addWidget(self.label_Schedule, 1, 0, 1, 1)
        self.lineEdit_Schedule = QtWidgets.QLineEdit(self.scrollAreaWidgetContents_2)
        font = QtGui.QFont()
        font.setPointSize(10)
        self.lineEdit_Schedule.setFont(font)
        self.lineEdit_Schedule.setObjectName("lineEdit_Schedule")
        self.gridLayout.addWidget(self.lineEdit_Schedule, 1, 1, 1, 1)
        self.pushButton_Schedule = QtWidgets.QPushButton(self.scrollAreaWidgetContents_2)
        self.pushButton_Schedule.setObjectName("pushButton_Schedule")
        self.gridLayout.addWidget(self.pushButton_Schedule, 1, 2, 1, 1)
        self.scrollArea_2.setWidget(self.scrollAreaWidgetContents_2)
        self.scrollArea_3 = QtWidgets.QScrollArea(self.centralwidget)
        self.scrollArea_3.setGeometry(QtCore.QRect(10, 160, 651, 52))
        self.scrollArea_3.setWidgetResizable(True)
        self.scrollArea_3.setObjectName("scrollArea_3")
        self.scrollAreaWidgetContents_3 = QtWidgets.QWidget()
        self.scrollAreaWidgetContents_3.setGeometry(QtCore.QRect(0, 0, 649, 50))
        self.scrollAreaWidgetContents_3.setObjectName("scrollAreaWidgetContents_3")
        self.gridLayout_3 = QtWidgets.QGridLayout(self.scrollAreaWidgetContents_3)
        self.gridLayout_3.setObjectName("gridLayout_3")
        self.dateEdit = QtWidgets.QDateEdit(self.scrollAreaWidgetContents_3)
        self.dateEdit.setDateTime(QtCore.QDateTime(QtCore.QDate(2020, 5, 1), QtCore.QTime(12, 0, 0)))
        self.dateEdit.setDate(QtCore.QDate(2020, 5, 1))
        self.dateEdit.setObjectName("dateEdit")
        self.gridLayout_3.addWidget(self.dateEdit, 0, 0, 1, 1)
        self.progressBar = QtWidgets.QProgressBar(self.scrollAreaWidgetContents_3)
        self.progressBar.setProperty("value", 0)
        self.progressBar.setObjectName("progressBar")
        self.gridLayout_3.addWidget(self.progressBar, 0, 1, 1, 1)
        self.pushButton_Labeling = QtWidgets.QPushButton(self.scrollAreaWidgetContents_3)
        self.pushButton_Labeling.setObjectName("pushButton_Labeling")
        self.gridLayout_3.addWidget(self.pushButton_Labeling, 0, 2, 1, 1)
        self.scrollArea_3.setWidget(self.scrollAreaWidgetContents_3)
        self.scrollArea_4 = QtWidgets.QScrollArea(self.centralwidget)
        self.scrollArea_4.setGeometry(QtCore.QRect(10, 220, 651, 291))
        self.scrollArea_4.setWidgetResizable(True)
        self.scrollArea_4.setObjectName("scrollArea_4")
        self.scrollAreaWidgetContents_4 = QtWidgets.QWidget()
        self.scrollAreaWidgetContents_4.setGeometry(QtCore.QRect(0, 0, 649, 289))
        self.scrollAreaWidgetContents_4.setObjectName("scrollAreaWidgetContents_4")
        self.listView = QtWidgets.QListView(self.scrollAreaWidgetContents_4)
        self.listView.setGeometry(QtCore.QRect(10, 10, 121, 261))
        self.listView.setObjectName("listView")
        self.listView2 = QtWidgets.QListView(self.scrollAreaWidgetContents_4)
        self.listView2.setGeometry(QtCore.QRect(140, 10, 501, 261))
        self.listView2.setObjectName("listView2")
        self.scrollArea_4.setWidget(self.scrollAreaWidgetContents_4)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 671, 25))
        self.menubar.setObjectName("menubar")
        self.menu = QtWidgets.QMenu(self.menubar)
        self.menu.setObjectName("menu")
        self.menu_2 = QtWidgets.QMenu(self.menubar)
        self.menu_2.setObjectName("menu_2")
        self.menu_4 = QtWidgets.QMenu(self.menubar)
        self.menu_4.setObjectName("menu_4")
        self.menu_5 = QtWidgets.QMenu(self.menubar)
        self.menu_5.setObjectName("menu_5")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)
        self.action_New = QtWidgets.QAction(MainWindow)
        self.action_New.setAutoRepeat(False)
        self.action_New.setObjectName("action_New")
        self.action_Open = QtWidgets.QAction(MainWindow)
        self.action_Open.setObjectName("action_Open")
        self.action_Window = QtWidgets.QAction(MainWindow)
        self.action_Window.setObjectName("action_Window")
        self.action_5 = QtWidgets.QAction(MainWindow)
        self.action_5.setObjectName("action_5")
        self.action_6 = QtWidgets.QAction(MainWindow)
        self.action_6.setObjectName("action_6")
        self.action = QtWidgets.QAction(MainWindow)
        self.action.setObjectName("action")
        self.action_2 = QtWidgets.QAction(MainWindow)
        self.action_2.setObjectName("action_2")
        self.menu.addAction(self.action_New)
        self.menu.addAction(self.action_Open)
        self.menu_2.addAction(self.action_Window)
        self.menu_5.addAction(self.action)
        self.menu_5.addAction(self.action_2)
        self.menubar.addAction(self.menu.menuAction())
        self.menubar.addAction(self.menu_5.menuAction())
        self.menubar.addAction(self.menu_2.menuAction())
        self.menubar.addAction(self.menu_4.menuAction())

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)
        MainWindow.setTabOrder(self.scrollArea, self.lineEdit_Excel)
        MainWindow.setTabOrder(self.lineEdit_Excel, self.scrollArea_2)
        MainWindow.setTabOrder(self.scrollArea_2, self.lineEdit_Serviceman)
        MainWindow.setTabOrder(self.lineEdit_Serviceman, self.pushButton_Serviceman)
        MainWindow.setTabOrder(self.pushButton_Serviceman, self.scrollArea_3)
        MainWindow.setTabOrder(self.scrollArea_3, self.lineEdit_Schedule)
        MainWindow.setTabOrder(self.lineEdit_Schedule, self.pushButton_Schedule)
        MainWindow.setTabOrder(self.pushButton_Schedule, self.dateEdit)
        MainWindow.setTabOrder(self.dateEdit, self.pushButton_Labeling)
        MainWindow.setTabOrder(self.pushButton_Labeling, self.scrollArea_4)

        # self.dateEdit.setDate(QtCore.QDate(1995, 5, 17))
        # Custom Code Here
        self.dir_name = os.path.join(os.environ['USERPROFILE'], 'Desktop')
        self.model = ExcelInstance(self.dateEdit.date().year(), self.dateEdit.date().month())
        self.update_listView2()

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.label_Excel.setText(_translate("MainWindow", "Excel 檔案"))
        self.label_Serviceman.setText(_translate("MainWindow", "役男輪值表"))
        self.pushButton_Serviceman.setText(_translate("MainWindow", "選擇檔案"))
        self.label_Schedule.setText(_translate("MainWindow", "排班表"))
        self.pushButton_Schedule.setText(_translate("MainWindow", "選擇檔案"))
        self.pushButton_Labeling.setText(_translate("MainWindow", "開始標記"))
        self.menu.setTitle(_translate("MainWindow", "檔案"))
        self.menu_2.setTitle(_translate("MainWindow", "測試"))
        self.menu_4.setTitle(_translate("MainWindow", "說明"))
        self.menu_5.setTitle(_translate("MainWindow", "修改"))
        self.action_New.setText(_translate("MainWindow", "新增檔案"))
        self.action_New.setShortcut(_translate("MainWindow", "Ctrl+N"))
        self.action_Open.setText(_translate("MainWindow", "開啟舊檔"))
        self.action_Open.setShortcut(_translate("MainWindow", "Ctrl+O"))
        self.action_Window.setText(_translate("MainWindow", "開新視窗"))
        self.action_Window.setShortcut(_translate("MainWindow", "Ctrl+W"))
        self.action_5.setText(_translate("MainWindow", "博愛苑"))
        self.action_6.setText(_translate("MainWindow", "仁愛苑"))
        self.action.setText(_translate("MainWindow", "教保科人員"))
        self.action_2.setText(_translate("MainWindow", "役男"))

        # Custom Code Here
        self.action_New.triggered.connect(self.excel_save)
        self.action_Open.triggered.connect(self.excel_open)
        self.action_Window.triggered.connect(self.Open_SelectedDialog)
        self.action.triggered.connect(lambda:self.Open_TxtFile("教保科人員.txt"))
        self.action_2.triggered.connect(lambda:self.Open_TxtFile("教保科役男.txt"))
        self.dateEdit.dateChanged.connect(self.onDateChanged)
        self.pushButton_Schedule.clicked.connect(self.Open_Folder)
        self.pushButton_Serviceman.clicked.connect(self.Open_Docx)
        self.pushButton_Labeling.clicked.connect(lambda: self.startlabel(MainWindow))

    def startlabel(self, MainWindow):
        # Pop a warning dialog to 
        if self.lineEdit_Excel.text() == '':
            self.Pop_WarnDialog("請先選擇Excel檔案")
            return

        if self.lineEdit_Schedule.text() == '':
            self.Pop_WarnDialog("請先選擇班表")
            return
            
        # if self.lineEdit_Serviceman.text() == '':
        #     msg = QtWidgets.QMessageBox()
        #     msg.setWindowTitle("Warning")
        #     msg.setText("請先選擇役男輪值表")
        #     msg.Icon(QtWidgets.QMessageBox.Warning)
        #     msg.setStandardButtons(QtWidgets.QMessageBox.Ok)
        #     x = msg.exec_()
        #     return


        if not self.model.startlabeling_1(self.lineEdit_Excel.text()):
            self.Pop_ErrorDialog("標記動作已取消")
            return

        global is_label
        is_label = True
        for i in self.model.startlabeling_2(self.lineEdit_Schedule.text()):
            self.progressBar.setProperty("value", self.progressBar.value()+ 1)
            self.update_listView2()

        if not is_label:
            self.Pop_ErrorDialog("標記動作已取消")
            return

        if not self.lineEdit_Serviceman.text() == '':
            if not self.model.startlabeling_3(self.lineEdit_Serviceman.text()):
                self.Pop_ErrorDialog("標記動作已取消")
                return
                
        self.progressBar.setProperty("value", 12)

        scale = 0
        for i in self.model.startlabeling_4():
            if i ==10:
                self.progressBar.setProperty("value", self.progressBar.value()+ 10)
            elif i:
                scale=i
            else:
                self.progressBar.setProperty("value", self.progressBar.value()+ (70/scale))

        self.model.reset_variable()
        msg = QtWidgets.QMessageBox()
        msg.setWindowTitle("Finish")
        msg.setText("標記完成")
        print("標記完成")
        msg.Icon(QtWidgets.QMessageBox.Information)
        msg.setStandardButtons(QtWidgets.QMessageBox.Ok)
        x = msg.exec_()
        MainWindow.close()

    def Open_TxtFile(self, filename):
        os.system(filename)
        
    def onDateChanged(self, qDate):
        self.model.set_Date(qDate.year(), qDate.month())
        print(f'{qDate.year()}/{qDate.month()}/{qDate.day()}')
        # print('{0}/{1}/{2}'.format(qDate.day(), qDate.month(), qDate.year()))

    def excel_save(self):
        name = QtWidgets.QFileDialog.getSaveFileName(MainWindow, 
                                                      "New File",
                                                      self.dir_name,
                                                      "Excel (*.xlsx)")
        if name[0]:
            self.dir_name = os.path.dirname(name[0])
            self.lineEdit_Excel.setText(name[0])
            # if os.path.exists(name[0]):
            #     os.remove(name[0])

    def excel_open(self):
        name = QtWidgets.QFileDialog.getOpenFileName(MainWindow,
                                                     "Open File",
                                                      self.dir_name,
                                                      "Excel (*.xlsx *.xls *.xlsx *.xlsm *.xlsb)")
        if name[0]:
            self.dir_name = os.path.dirname(name[0])
            self.lineEdit_Excel.setText(name[0])

    # 選擇役男輪值表
    def Open_Docx(self):
        if self.lineEdit_Excel.text() == '':
            self.Pop_WarnDialog("請先選擇一個Excel檔案")
            return

        name = QtWidgets.QFileDialog.getOpenFileName( MainWindow,
                                                      "Open File",
                                                      self.dir_name,
                                                      "Word (*.docx)")
        print(name)
        # if the direction exists, added to lineEdit.
        if name[0]:
            self.dir_name = os.path.dirname(name[0])
            self.lineEdit_Serviceman.setText(name[0])
        return

    # 選擇排班表
    def Open_Folder(self):
        # if self.lineEdit_Excel.text() == '':
        #     self.Pop_WarnDialog("請先選擇一個Excel檔案")
        #     return

        folder = QtWidgets.QFileDialog.getExistingDirectory( MainWindow,
                                                             "Open Directory",
                                                             self.dir_name)
        if folder:
            # Try to set year and month according to the folder name.
            try:
                [year, month] = os.path.basename(folder).split('.')
                self.dateEdit.setDate(QtCore.QDate( int(year)+1911, int(month), 1))
                self.model.set_Date( int(year)+1911, int(month))

            except:
                self.Pop_ErrorDialog("格式錯誤，請手動設定班表月份")

            files = [os.path.basename(i) for i in glob.glob(os.path.join(folder, '*.docx'))]
            slm = QStringListModel()
            slm.setStringList(files)
            self.listView.setModel(slm)
            del slm

            if self.model.load_employees(folder):
                self.lineEdit_Schedule.setText(folder)
            self.update_listView2()
        return

    def update_listView2(self):
        self.model.load_employee_EDepart()
        Dict_Member = self.model.get_employees()
        list_employees = []
        for depart in Dict_Member:
            if depart=='教':
                list_employees.append('教保科(員): '+', '.join(Dict_Member[depart][0]))
                list_employees.append('教保科(役男): '+', '.join(Dict_Member[depart][1]))
            else:                
                list_employees.append(depart+'愛院(員): '+', '.join(Dict_Member[depart][0]))
                list_employees.append(depart+'愛院(工): '+', '.join(Dict_Member[depart][1]))
        slm = QStringListModel()
        slm.setStringList(list_employees)
        self.listView2.setModel(slm)
        del slm

    def Pop_WarnDialog(self, message):
        msg = QtWidgets.QMessageBox()
        msg.setWindowTitle("Warning")
        msg.setText(message)
        print(message)
        msg.Icon(QtWidgets.QMessageBox.Warning)
        msg.setStandardButtons(QtWidgets.QMessageBox.Ok)
        msg.exec_()
        del msg
        return

    def Pop_ErrorDialog(self, message):
        self.model.reset_variable()
        self.progressBar.setProperty("value", 0)
        msg = QtWidgets.QMessageBox()
        msg.setWindowTitle("Cancel")
        msg.setText(message)
        print(message)
        msg.Icon(QtWidgets.QMessageBox.Information)
        msg.setStandardButtons(QtWidgets.QMessageBox.Ok)
        msg.exec_()
        del msg
        return

    def Open_SelectedDialog(self):
        mList = ['1','2','3','4','5']
        Window = QtWidgets.QDialog()
        ui = SelectDialog()
        ui.setupUi(Window, mList.copy(), mList.copy(), "就你最特別")
        Window.setModal(True)
        Window.show()
        if Window.exec_():
            print(ui.getState())
        return

class SelectDialog(object):
    def setupUi(self, Form_SelectedWindow, List1, List2, labeltext):
        self.select = ['','']
        Form_SelectedWindow.setObjectName("Form_SelectedWindow")
        Form_SelectedWindow.resize(480, 480)
        self.gridLayout = QtWidgets.QGridLayout(Form_SelectedWindow)
        self.gridLayout.setObjectName("gridLayout")
        self.scrollArea = QtWidgets.QScrollArea(Form_SelectedWindow)
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setObjectName("scrollArea")
        self.scrollAreaWidgetContents = QtWidgets.QWidget()
        self.scrollAreaWidgetContents.setGeometry(QtCore.QRect(0, 0, 341, 434))
        self.scrollAreaWidgetContents.setObjectName("scrollAreaWidgetContents")
        self.gridLayout_2 = QtWidgets.QGridLayout(self.scrollAreaWidgetContents)
        self.gridLayout_2.setObjectName("gridLayout_2")
        self.pushButton_OK = QtWidgets.QPushButton(self.scrollAreaWidgetContents)
        font = QtGui.QFont()
        font.setFamily("微軟正黑體")
        self.pushButton_OK.setFont(font)
        self.pushButton_OK.setObjectName("pushButton_OK")
        self.gridLayout_2.addWidget(self.pushButton_OK, 3, 0, 1, 1)
        self.tabWidget = QtWidgets.QTabWidget(self.scrollAreaWidgetContents)
        self.tabWidget.setObjectName("tabWidget")
        self.tab_3 = QtWidgets.QWidget()
        self.tab_3.setObjectName("tab_3")
        self.gridLayout_3 = QtWidgets.QGridLayout(self.tab_3)
        self.gridLayout_3.setObjectName("gridLayout_3")
        self.listView = QtWidgets.QListView(self.tab_3)
        font = QtGui.QFont()
        font.setFamily("微軟正黑體")
        font.setPointSize(12)
        self.listView.setFont(font)
        self.listView.setObjectName("listView")
        self.gridLayout_3.addWidget(self.listView, 0, 0, 1, 1)
        self.tabWidget.addTab(self.tab_3, "")
        self.tab_4 = QtWidgets.QWidget()
        self.tab_4.setObjectName("tab_4")
        self.gridLayout_4 = QtWidgets.QGridLayout(self.tab_4)
        self.gridLayout_4.setObjectName("gridLayout_4")
        self.listView2 = QtWidgets.QListView(self.tab_4)
        font = QtGui.QFont()
        font.setFamily("微軟正黑體")
        font.setPointSize(12)
        self.listView2.setFont(font)
        self.listView2.setObjectName("listView2")
        self.gridLayout_4.addWidget(self.listView2, 0, 0, 1, 1)
        self.tabWidget.addTab(self.tab_4, "")
        self.gridLayout_2.addWidget(self.tabWidget, 1, 0, 1, 1)
        self.label = QtWidgets.QLabel(self.scrollAreaWidgetContents)
        font = QtGui.QFont()
        font.setFamily("微軟正黑體")
        font.setPointSize(12)
        self.label.setFont(font)
        self.label.setObjectName("label")
        self.gridLayout_2.addWidget(self.label, 0, 0, 1, 1)
        self.scrollArea.setWidget(self.scrollAreaWidgetContents)
        self.gridLayout.addWidget(self.scrollArea, 0, 0, 1, 1)

        # Custom Code Here
        self.List1 = List1
        self.List1.extend(["新增員","新增教保科員","忽略"])
        slm = QStringListModel()
        slm.setStringList(self.List1)
        self.listView.setModel(slm)
        del slm

        self.List2 = List2
        self.List2.extend(["新增工","新增教保科員","忽略"])
        slm = QStringListModel()
        slm.setStringList(self.List2)
        self.listView2.setModel(slm)
        del slm
        self.retranslateUi(Form_SelectedWindow, labeltext)
        self.tabWidget.setCurrentIndex(0)
        QtCore.QMetaObject.connectSlotsByName(Form_SelectedWindow)

    def retranslateUi(self, Form_SelectedWindow, labeltext):
        _translate = QtCore.QCoreApplication.translate
        Form_SelectedWindow.setWindowTitle(_translate("Form_SelectedWindow", "人員選擇"))
        self.pushButton_OK.setText(_translate("Form_SelectedWindow", "確定"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab_3), _translate("Form_SelectedWindow", "員"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab_4), _translate("Form_SelectedWindow", "工"))
        # self.label.setText(_translate("Form_SelectedWindow", "在*院找不到員工 小玠  (就你最特別)"))
        self.label.setText(_translate("Form_SelectedWindow", labeltext))
        self.listView.clicked.connect(self.clickedlist)
        self.listView2.clicked.connect(self.clickedlist2)
        self.pushButton_OK.clicked.connect(lambda: self.Window_Close(Form_SelectedWindow))
        # self.pushButton_OK.clicked.connect(Form_SelectedWindow.accept)

    def clickedlist(self, qModelIndex):
        self.select[0] = self.List1[qModelIndex.row()]
        # print("點擊的是：" + self.select[0])

    def clickedlist2(self, qModelIndex):
        self.select[1] = self.List2[qModelIndex.row()]
        # print("點擊的是：" + self.select[1])

    def Window_Close(self, Form_SelectedWindow):
        if not self.select[self.tabWidget.currentIndex()]:
            msg = QtWidgets.QMessageBox()
            msg.setWindowTitle("Error")
            msg.setText("請先選擇一個選項")
            msg.Icon(QtWidgets.QMessageBox.Warning)
            msg.setStandardButtons(QtWidgets.QMessageBox.Ok)
            x = msg.exec_()
        else:
            Form_SelectedWindow.accept()

    def getState(self):
        return self.select[self.tabWidget.currentIndex()]

class CalendarDialog(object):
    def setupUi(self, Form, Year, Month):
        self.Day, self.Dates = monthrange(Year, Month)
        Form.setObjectName("Form")
        Form.resize(410, 321)
        self.gridLayout = QtWidgets.QGridLayout(Form)
        self.gridLayout.setObjectName("gridLayout")
        self.calendarWidget = CalendarWidget(Form)
        self.calendarWidget.setupUi(Year, Month)
        self.calendarWidget.setCursor(QtGui.QCursor(QtCore.Qt.ArrowCursor))
        self.calendarWidget.setMouseTracking(False)
        self.calendarWidget.setTabletTracking(False)
        self.calendarWidget.setSelectedDate(QtCore.QDate(Year, Month, 1))
        self.calendarWidget.setMinimumDate(QtCore.QDate(Year, Month, 1))
        self.calendarWidget.setMaximumDate(QtCore.QDate(Year, Month, self.Dates))
        self.calendarWidget.setGridVisible(True)
        self.calendarWidget.setHorizontalHeaderFormat(QtWidgets.QCalendarWidget.ShortDayNames)
        self.calendarWidget.setNavigationBarVisible(True)
        self.calendarWidget.setDateEditEnabled(False)
        self.calendarWidget.setObjectName("calendarWidget")
        self.gridLayout.addWidget(self.calendarWidget, 0, 0, 1, 1)
        self.pushButton_OK = QtWidgets.QPushButton(Form)
        self.pushButton_OK.setObjectName("pushButton_OK")
        self.gridLayout.addWidget(self.pushButton_OK, 1, 0, 1, 1)

        self.retranslateUi(Form)
        QtCore.QMetaObject.connectSlotsByName(Form)

    def retranslateUi(self, Form):
        _translate = QtCore.QCoreApplication.translate
        Form.setWindowTitle(_translate("Form", "Form"))
        self.pushButton_OK.setText(_translate("Form", "OK"))
        self.calendarWidget.clicked.connect(self.inverse_state)
        self.pushButton_OK.clicked.connect(lambda: self.Window_Close(Form))

    def inverse_state(self, qDate):
        self.calendarWidget.inverse_state(qDate.day())
        self.calendarWidget.updateCells()

    def Window_Close(self, Form):
        Form.accept()

    def get_Holiday(self):
        return self.calendarWidget.get_Holiday()

class CalendarWidget(QtWidgets.QCalendarWidget):
    def __init__(self, parent=None):
        super(CalendarWidget, self).__init__(parent,
            verticalHeaderFormat=QtWidgets.QCalendarWidget.NoVerticalHeader,
            gridVisible=True)

        self.color = QtGui.QColor(0,255,0,64)

        for d in (QtCore.Qt.Saturday, QtCore.Qt.Sunday,):
            fmt = self.weekdayTextFormat(d)
            fmt.setForeground(QtCore.Qt.darkGray)
            self.setWeekdayTextFormat(d, fmt)
        # self.selectionChanged.connect(self.updateCells)

    def setupUi(self, Year, Month):
        self.Day, self.Dates = monthrange(Year, Month)
        # set Holiday
        self.Holiday = {}
        for d in range(self.Dates):
            if (self.Day+d)%7 ==5 or (self.Day+d)%7 ==6:
                self.Holiday[d+1]=True
            else:
                self.Holiday[d+1]=False

    def inverse_state(self, date):
        if self.Holiday[date]:
            self.Holiday[date] = False
        else:
            self.Holiday[date] = True

    def get_Holiday(self):
        return self.Holiday

    def paintCell(self, painter, rect, date):
        if (date < self.minimumDate()) or (date > self.maximumDate()):
            super(CalendarWidget, self).paintCell(painter, rect, date)
        elif self.Holiday[date.day()]:
            # painter.fillRect(rect, QtGui.QColor("#D3D3D3"))
            painter.fillRect(rect, self.color)
            painter.drawText(rect, QtCore.Qt.AlignCenter, str(date.day()))
        else:
            painter.drawText(rect, QtCore.Qt.AlignCenter, str(date.day()))

class ExcelInstance(object):
    def __init__(self, Year, Month):
        self.Table_Width= 0
        self.Table_Height = 22   # (10個苑+教保科)*列
        self.Table_Index = []
        self.Dict_Member = {}
        self.Dict_Nickname = {}
        self.List_Ignore = []
        self.Year = Year
        self.Month = Month
        self.Day, self.Dates = monthrange(Year, Month)

        for depart, member in Departments.__members__.items():
            if depart=='教':
                self.Table_Index.append([depart, "員"])
                self.Table_Index.append([depart, "役男"])
                continue

            self.Table_Index.append([depart, "員"])
            self.Table_Index.append([depart, "工"])

    def set_Date(self,Year, Month):
        self.Year = Year
        self.Month = Month
        self.Day, self.Dates = monthrange(Year, Month)

    def load_employee_EDepart(self):
        depart='教'
        member_list = []
        with open(os.path.join(os.getcwd(), "教保科人員.txt"), 'r', encoding='utf-8') as f:
            temp = []
            for line in f.read().split('\n'):
                for name in line.split(','):
                    temp.append(name.rstrip())
            member_list.append(temp)
            
        with open(os.path.join(os.getcwd(), "教保科役男.txt"), 'r', encoding='utf-8') as f:
            temp = []
            for line in f.read().split('\n'):
                for name in line.split(','):
                    temp.append(name.rstrip())
            member_list.append(temp)
        self.Dict_Member[depart] = member_list
        # self.Table_Height = self.Table_Height + np.array(pd.DataFrame(member_list)).shape[0]
        if np.array(pd.DataFrame(member_list)).shape[1] > self.Table_Width:
            self.Table_Width = np.array(pd.DataFrame(member_list)).shape[1]
        return

    def load_employees(self, dirname):
        for depart, member in Departments.__members__.items():
            if depart=='教':
                continue           
            # load employees list from excel
            try:
                # df = pd.read_excel(glob.glob(os.path.join(dirname, '*'+depart+'*.xlsx'))[0], index=None, sheet_name=-1)
                ExcelFile = pd.ExcelFile(glob.glob(os.path.join(dirname, '*'+depart+'*.xlsx'))[0])
                for sheet_name in ExcelFile.sheet_names[::-1]:
                    if not '工作表' in sheet_name:
                        df = ExcelFile.parse(sheet_name)
                        break

                # reset the head of DataFrame
                columns = [col.strip() for col in df.columns]
                columns[0]=''; columns[1]='員工'; columns[2]=''
                df.columns = columns
                check_repeat = [name.rstrip()[-2:] for idx, name in zip(df['員工'],df['姓名']) if not (idx is np.nan)]
            except:
                msg = QtWidgets.QMessageBox()
                msg.setWindowTitle("Cancel")
                msg.setText("請檢察 1."+ depart+"愛苑人員統計表是否存在，2.格式錯誤 3.檔案是否處於開啟狀態")
                print("請檢察 1."+ depart+"愛苑人員統計表是否存在，2.格式錯誤 3.檔案是否處於開啟狀態")
                msg.Icon(QtWidgets.QMessageBox.Information)
                msg.setStandardButtons(QtWidgets.QMessageBox.Ok)
                x = msg.exec_()
                return False

            member_list = []
            temp = []
            for name in df[df['員工']=='督']['姓名']:
                if check_repeat.count(name.rstrip()[-2:])>1:
                    temp.append(name.rstrip())
                else:
                    temp.append(name[-2:])
                    
            for name in df[df['員工']=='員']['姓名']:
                if check_repeat.count(name.rstrip()[-2:])>1:
                    temp.append(name.rstrip())
                else:
                    temp.append(name[-2:])
            member_list.append(temp)
            
            temp = []
            for name in df[df['員工']=='工']['姓名']:
                if check_repeat.count(name.rstrip()[-2:])>1:
                    temp.append(name.rstrip())
                else:
                    temp.append(name.rstrip()[-2:])
            member_list.append(temp)
            
            # self.Table_Height = self.Table_Height+ np.array(pd.DataFrame(member_list)).shape[0]
            if np.array(pd.DataFrame(member_list)).shape[1] > self.Table_Width:
                self.Table_Width = np.array(pd.DataFrame(member_list)).shape[1]
        #     List_Member[depart] = np.array(pd.DataFrame(member_list))
            self.Dict_Member[depart] = member_list
            del member_list, temp

        self.Schedule_First = np.ndarray((self.Dates, len(Departments), 2, self.Table_Width), dtype='O')
        self.Schedule_Second_Name = np.ndarray((self.Dates, len(Departments), 2, 1), dtype='O')
        self.Schedule_Third_Name = [[[  [] for x in range(2)] for y in range(len(Departments))] for k in range(self.Dates)]
        self.Schedule_Third_State = [[[  [] for x in range(2)] for y in range(len(Departments))] for k in range(self.Dates)]
        return True

    def reset_variable(self):
        self.Schedule_First = np.ndarray((self.Dates, len(Departments), 2, self.Table_Width), dtype='O')
        self.Schedule_Second_Name = np.ndarray((self.Dates, len(Departments), 2, 1), dtype='O')
        self.Schedule_Third_Name = [[[  [] for x in range(2)] for y in range(len(Departments))] for k in range(self.Dates)]
        self.Schedule_Third_State = [[[  [] for x in range(2)] for y in range(len(Departments))] for k in range(self.Dates)]
 
    def get_employees(self):
        return self.Dict_Member

    def iloc_search(self, date, depart, name, state, add_timeoff_before= False, add_middle_before= False):

        List_Member = np.array(pd.DataFrame(self.Dict_Member[depart]))
        
        if name[-2:] in self.List_Ignore:
            return True
        # Switch Nickname if exist.
        if name[-2:] in self.Dict_Nickname.keys():
            name = self.Dict_Nickname[name[-2:]]
        # Start find the iloc
        if np.size(np.argwhere( List_Member == name[-2:])) > 0:
            i, j = np.argwhere(List_Member == name[-2:])[0]
            
            if add_timeoff_before:
                self.Schedule_Third_Name[date-1][Departments[depart].value][i].append(name[-2:]+'('+state+')')
                state = self.handle_state(state)
                self.Schedule_Third_State[date-1][Departments[depart].value][i].append(state)
                
            if add_middle_before:
                if self.Schedule_First[date-1, Departments[depart].value, i, j] == '中班B':
                    self.Schedule_Second_Name[date-1, Departments[depart].value, i] = name[-2:]
                    
            self.Schedule_First[date-1, Departments[depart].value, i, j] = state
            
            del i, j
            return True
            
        elif np.size(np.argwhere( List_Member == name[-3:])) > 0:
            i, j = np.argwhere(List_Member == name[-3:])[0]
            
            if add_timeoff_before:
                self.Schedule_Third_Name[date-1][Departments[depart].value][i].append(name[-3:]+'('+state+')')
                state = self.handle_state(state)
                self.Schedule_Third_State[date-1][Departments[depart].value][i].append(state)

            if add_middle_before:
                if self.Schedule_First[date-1, Departments[depart].value, i, j] == '中班B':
                    self.Schedule_Second_Name[date-1, Departments[depart].value, i] = name[-3:]
                    
            self.Schedule_First[date-1, Departments[depart].value, i, j] = state
            del i, j
            return True
        
        # Find iloc in deferent department
        elif depart != '教':
            if self.iloc_search(date, '教', name, state, add_timeoff_before):
                return True
        else:
            return False
        
        print("{0}愛院 Day{1}, {2} 找不到人員{3}".format(depart, date, state, name))

        Window = QtWidgets.QDialog()
        ui = SelectDialog()
        ui.setupUi(Window, self.Dict_Member[depart][0].copy(), 
                           self.Dict_Member[depart][1].copy(),
                           f'{depart}愛院班表(日期:{date},班別:{state})中, 找不到 \"{name}\"')
        Window.setModal(True)
        Window.show()
        if Window.exec_() == QtWidgets.QDialog.Accepted:
            print("點擊的是：" + ui.getState())
            if ui.getState()=='忽略':
                self.List_Ignore.append(name[-2:])
            elif ui.getState()=='新增員':
                self.add_employee(depart, name[-2:], 0)
            elif ui.getState()=='新增工':
                self.add_employee(depart, name[-2:], 1)
            elif ui.getState()=='新增教保科員':
                self.add_employee('教', name[-2:], 0)
            else:
                self.Dict_Nickname[name[-2:]] = ui.getState()
            return self.iloc_search(date, depart, name, state, add_timeoff_before)
        else:
            # Stop the labeling when user cancel the dialog.
            global is_label
            is_label = False
            return False
        return False

    # Handle the time off text
    def split_text(self, text):
        temp = []    
        text = text.replace('（','(').replace('）',')').replace(':','')
        for line in text.rstrip().split('\n'):
            for name in line.split('.'):
                if len(name.rstrip())>0:
                    if ' ' in name.rstrip():
                        temp.extend([i for i in name.rstrip().split(' ') if len(i)>0])
                    else:
                        temp.append(name.rstrip())
                        
        for idx, name in enumerate(temp):
            if name.find(')')>0 and (name.find(')')+1) != len(name):
                temp[idx]= name[:name.find(')')+1]
                temp.insert( idx+1, name[name.find(')')+1:])
        for idx, name in enumerate(temp):
            if name[-1]=='(':
                temp[idx:idx+2]= [''.join(temp[idx:idx+2])]
                
        return temp

    # Handle the time off state
    def handle_state(self, state):
        if state.rstrip() =='上午':
            return '上休'
        
        elif state.rstrip() =='上':
            return '上休'
        
        elif state.rstrip() == '下午':
            return '下休'
        
        elif state.rstrip() == '下':
            return '下休'
        
        else:
            a, b = state.split('-')
            if int(a.rstrip()[:2]) > 12:
                return '下休'
            
            elif int(b.rstrip()[:2]) > 12:
                return '輪休'
            
            else:
                return '上休'

    def timeoff(self, depart, row_cells):
        date = int(row_cells[0].text.rstrip())
        names = self.split_text(row_cells[12].text)
        for name in names:
            if not is_label:
                return
            if '(' in name:
                name, state = name.split('(')
                self.iloc_search(date, depart, name.rstrip(), state[:-1].rstrip(), add_timeoff_before = True)
                
            else:   
                self.iloc_search(date, depart, name.rstrip(), "輪休")

    def night(self, depart, row_cells):
        date = int(row_cells[0].text.rstrip())
        
        if row_cells[10].text.rstrip():
            self.iloc_search(date, depart, row_cells[10].text.rstrip(), "夜班")
            
        if row_cells[9].text.rstrip():
            self.iloc_search(date, depart, row_cells[9].text.rstrip(), "夜班")
            
        if row_cells[8].text.rstrip():
            self.iloc_search(date, depart, row_cells[8].text.rstrip(), "夜班")

    def middle(self, depart, row_cells):
        date = int(row_cells[0].text.rstrip())
        
        if row_cells[7].text.rstrip():
            self.iloc_search(date, depart, row_cells[7].text.rstrip(), "中班")
            
        if row_cells[6].text.rstrip():
            self.iloc_search(date, depart, row_cells[6].text.rstrip(), "中班B")
            
        if row_cells[5].text.rstrip():
            self.iloc_search(date, depart, row_cells[5].text.rstrip(), "中班")

    def morning(self, depart, row_cells):
        date = int(row_cells[0].text.rstrip())
        
        if row_cells[4].text.rstrip():
            self.iloc_search(date, depart, row_cells[4].text.rstrip(), "早班", add_middle_before = True)
            
        if row_cells[3].text.rstrip():
            self.iloc_search(date, depart, row_cells[3].text.rstrip(), "中班")
            
        if row_cells[2].text.rstrip():
            # if (self.Day+date-1)%7 ==5 or (self.Day+date-1)%7 ==6:
            if self.Holiday[date]:
                self.iloc_search(date, depart, row_cells[2].text.rstrip(), "假日班")
            else:
                self.iloc_search(date, depart, row_cells[2].text.rstrip(), "早班")

    def add_employee(self, depart, name, workstate):      
        self.Dict_Member[depart][workstate].append(name)  
        if np.array( pd.DataFrame( self.Dict_Member[depart])).shape[1] > self.Table_Width:
            self.Table_Width = np.array( pd.DataFrame( self.Dict_Member[depart])).shape[1]
            np.insert(self.Schedule_First, self.Schedule_First.shape[3], None, axis=3)


    def startlabeling_1(self, excel_name, cover = False):
        #開啟excel
        if cover and os.path.exists(excel_name):
            os.remove(excel_name)

        if not os.path.exists(excel_name):
            writer = pd.ExcelWriter(excel_name, engine='xlsxwriter')
            writer.save()
            writer.close()

        wb = xw.Book(excel_name)
        
        try:
            self.ws = wb.sheets.add(str(self.Year-1911)+'-'+str(self.Month), after= wb.sheets[-1])
            # self.ws = wb.sheets[str(self.Year-1911)+'-'+str(self.Month)]
        except:
            # self.ws = wb.sheets.add(str(self.Year-1911)+'-'+str(self.Month), after= wb.sheets[-1])
            msg = QtWidgets.QMessageBox()
            msg.setWindowTitle("Error")
            msg.setText("{0}名稱工作表已存在，是否要覆蓋".format(str(self.Year-1911)+'-'+str(self.Month)))
            msg.Icon(QtWidgets.QMessageBox.Warning)
            msg.setStandardButtons(QtWidgets.QMessageBox.Ok|QtWidgets.QMessageBox.Cancel)
            if msg.exec_() == QtWidgets.QMessageBox.Ok:
                wb.sheets[str(self.Year-1911)+'-'+str(self.Month)].delete()
                self.ws = wb.sheets.add(str(self.Year-1911)+'-'+str(self.Month), after= wb.sheets[-1])
            else:
                return False

        Window = QtWidgets.QDialog()
        ui = CalendarDialog()
        ui.setupUi(Window, self.Year, self.Month)
        Window.setModal(True)
        Window.show()
        self.Holiday = {}
        if Window.exec_():
            self.Holiday = ui.get_Holiday()
        else:
            return False

        # 教保科假日輪休
        xs, ys = np.where( np.array( pd.DataFrame(self.Dict_Member["教"]))!=None)
        for d in range(self.Dates):
            # if (self.Day+d)%7 ==5 or (self.Day+d)%7 ==6:
            if self.Holiday[d+1]:
                for i, j in zip(xs, ys):
                    self.Schedule_First[d, Departments["教"].value, i, j] = "輪休"
        return True

    def startlabeling_2(self, dirname):
        # 各苑排班表
        for depart, member in Departments.__members__.items():
            if depart=='教':
                continue
            try:
                print("loading: ", glob.glob(os.path.join(dirname, '*'+depart+'*.docx'))[0])
                document = Document(glob.glob(os.path.join(dirname, '*'+depart+'*.docx'))[0])
            except:
                global is_label
                is_label = False
                msg = QtWidgets.QMessageBox()
                msg.setWindowTitle("Cancel")
                msg.setText("找不到"+ dapart+"愛苑排班表")
                print("找不到"+ dapart+"愛苑排班表")
                msg.Icon(QtWidgets.QMessageBox.Information)
                msg.setStandardButtons(QtWidgets.QMessageBox.Ok)
                x = msg.exec_()
                return

            tables = document.tables
            for table in tables:
                for i in range(len(table.rows)):
                    if table.row_cells(i)[0].text.rstrip().isdigit():
                        # 1. 處理輪休
                        self.timeoff(depart, table.row_cells(i))
                        if not is_label:
                            return
                        # 2. 處理夜班
                        self.night(depart, table.row_cells(i))
                        if not is_label:
                            return
                        # 3. 處理中班
                        self.middle(depart, table.row_cells(i))
                        if not is_label:
                            return
                        # 4. 處理早班
                        self.morning(depart, table.row_cells(i))
                        if not is_label:
                            return
            yield
        return

    def startlabeling_3(self, docx_name):
        # 役男輪值表
        print(docx_name)
        document = Document(docx_name)
        print(document)
        table = document.tables[0]
        try:
            for i in range(len(table.rows)):
                if table.row_cells(i)[0].text.rstrip().isdigit():
                    date = int(table.row_cells(i)[0].text.rstrip())        
                    names = [j.split('、') for j in table.row_cells(i)[8].text.rstrip().split('\n')]
                    names = sum( names, [])
                    for name in names:
                        if not name:
                            continue
                        if '(' in name:
                            name, state = name.split('(')
                            state = self.handle_state(state[:-1])
                            self.iloc_search(date, "教", name.rstrip(), state)
                        else:
                            self.iloc_search(date, "教", name.rstrip(), "輪休")
                    name = table.row_cells(i)[2].text.rstrip()
                    self.iloc_search(date, "教", name.rstrip(), "早班")
                    name = table.row_cells(i)[3].text.rstrip()
                    self.iloc_search(date, "教", name.rstrip(), "備勤")
        except:
            msg = QtWidgets.QMessageBox()
            msg.setWindowTitle("Cancel")
            msg.setText("役男輪值表包含錯誤格式")
            print("役男輪值表包含錯誤格式")
            msg.Icon(QtWidgets.QMessageBox.Information)
            msg.setStandardButtons(QtWidgets.QMessageBox.Ok)
            x = msg.exec_()
            return False
        return True

    def startlabeling_4(self):
        Table_1 = []
        for depart, member in Departments.__members__.items():
            for i in self.Dict_Member[depart]:
                Table_1.append(i)
        Table_1 = np.array( pd.DataFrame(Table_1))
        Table_2 = self.Schedule_Second_Name.reshape(-1,1)
        Table_3 = np.array(pd.DataFrame(np.array(self.Schedule_Third_Name).reshape(-1).tolist()))

        self.ws[0,0].value = head_1
        self.ws[0, 3+ Table_1.shape[1]+ 3+ Table_3.shape[1]+3].value = head_2
        for i in range(self.Dates):
            self.ws[self.Table_Height*i+1: self.Table_Height*(i+1)+1, 0].value = i+1
            self.ws[self.Table_Height*i+1: self.Table_Height*(i+1)+1, 3+ Table_1.shape[1]+ 3+ Table_3.shape[1]+3].value = i+1
            
            self.ws[self.Table_Height*i+1: self.Table_Height*(i+1)+1, 1:3].value = self.Table_Index
            self.ws[self.Table_Height*i+1: self.Table_Height*(i+1)+1, 3].value = Table_1
            
            # if (self.Day+i)%7 ==5 or (self.Day+i)%7 ==6:
            if self.Holiday[i+1]:
                self.ws[self.Table_Height*i+1: self.Table_Height*(i+1)+1, 0].color = Color.輪休.value
                self.ws[self.Table_Height*i+1: self.Table_Height*(i+1)+1, 3+ Table_1.shape[1]+ 3+ Table_3.shape[1]+3].color = Color.輪休.value
                
        self.ws[1, 3+ Table_1.shape[1]+ 2].value = Table_2
        self.ws[1, 3+ Table_1.shape[1]+ 3].value = Table_3
        self.Schedule_First = self.Schedule_First.reshape(-1, self.Schedule_First.shape[-1])
        xs, ys = np.where( self.Schedule_First != None)

        yield (len(xs)//100)
        count_yield = 0
        for i, j in zip(xs, ys):
            if self.Schedule_First[i][j]=='上休':
                self.ws[ int(i+1), int(j+3)].api.Font.Color = Color.上休.value
                self.ws[ int(i+1), int(j+3)].color = Color.輪休.value
            elif self.Schedule_First[i][j]=='下休':
                self.ws[ int(i+1), int(j+3)].api.Font.Color = Color.下休.value
                self.ws[ int(i+1), int(j+3)].color = Color.輪休.value        
            else:
                self.ws[ int(i+1), int(j+3)].color = Color[ self.Schedule_First[i,j]].value

            if (count_yield%100)==0:
                yield
            count_yield = count_yield+1

        # Second Part
        self.Schedule_Second_Name = self.Schedule_Second_Name.reshape(-1, self.Schedule_Second_Name.shape[-1])
        xs, ys = np.where( self.Schedule_Second_Name != None)
        for i in xs:
            self.ws[ int(i+1), 3+ Table_1.shape[1]+ 2].color = Color.中班B.value
        yield 10

        # Third Part
        self.Third_State = np.array(pd.DataFrame(np.array(self.Schedule_Third_State).reshape(-1).tolist()))
        xs, ys = np.where( self.Third_State != None)
        for i, j in zip(xs, ys):
            if self.Third_State[i][j]=='上休':
                self.ws[ int(i+1), int(j+ 3+ Table_1.shape[1]+ 3)].api.Font.Color = Color.上休.value
                self.ws[ int(i+1), int(j+ 3+ Table_1.shape[1]+ 3)].color = Color.輪休.value
            elif self.Third_State[i][j]=='下休':
                self.ws[ int(i+1), int(j+ 3+ Table_1.shape[1]+ 3)].api.Font.Color = Color.下休.value
                self.ws[ int(i+1), int(j+ 3+ Table_1.shape[1]+ 3)].color = Color.輪休.value        
            else:
                self.ws[ int(i+1), int(j+ 3+ Table_1.shape[1]+ 3)].color = Color[ self.Third_State[i,j]].value
        yield 10

if __name__ == "__main__":
    appctxt = ApplicationContext()       # 1. Instantiate ApplicationContext
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    exit_code = appctxt.app.exec_()      # 2. Invoke appctxt.app.exec_()
    sys.exit(exit_code)